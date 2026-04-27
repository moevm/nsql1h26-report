import hashlib
import re
from io import BytesIO
from typing import List, Dict, Any

from docx import Document


VOWELS = set("аеёиоуыьъэюяАЕЁИОУЫЬЪЭЮЯaeiouAEIOU")


def count_syllables(word: str) -> int:
    return max(1, sum(1 for c in word if c in VOWELS))


def split_sentences(text: str) -> List[str]:
    sentences = re.split(r'(?<=[.!?])\s+', text.strip())
    return [s.strip() for s in sentences if s.strip()]


def split_into_chunks(text: str, min_len: int = 40) -> List[Dict[str, str]]:
    sentences = split_sentences(text)
    chunks = []
    buffer = ""
    for sent in sentences:
        buffer = (buffer + " " + sent).strip() if buffer else sent
        if len(buffer) >= min_len:
            chunks.append(buffer)
            buffer = ""
    if buffer:
        chunks.append(buffer)

    return [
        {"text": c, "hash": hashlib.sha256(c.encode("utf-8")).hexdigest()}
        for c in chunks
        if c.strip()
    ]


def compute_flesh_index(text: str) -> int:
    words = text.split()
    sentences = split_sentences(text)
    if not words or not sentences:
        return 0
    total_syllables = sum(count_syllables(w) for w in words)
    avg_sentence_len = len(words) / len(sentences)
    avg_syllables = total_syllables / len(words)
    score = 206.835 - 1.015 * avg_sentence_len - 84.6 * avg_syllables
    return max(0, min(100, int(score)))


def compute_keyword_density(text: str, top_n: int = 5) -> int:
    words = re.findall(r'\b[а-яёА-ЯЁa-zA-Z]{4,}\b', text.lower())
    if not words:
        return 0
    freq: Dict[str, int] = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    top = sorted(freq.values(), reverse=True)[:top_n]
    return min(100, int(sum(top) / len(words) * 100))


SECTION_KEYWORDS = {
    "введение": ["введен", "вводн"],
    "заключение": ["заключен", "вывод"],
    "список литературы": ["литератур", "библиограф", "список использ"],
    "основная часть": ["основн", "теорет", "практич", "реализ", "описан"],
}


def detect_section_type(heading: str) -> str:
    h = heading.lower()
    for stype, keywords in SECTION_KEYWORDS.items():
        if any(kw in h for kw in keywords):
            return stype
    return heading.strip() or "раздел"


def is_heading(para) -> bool:
    style_name = para.style.name.lower() if para.style else ""
    return (
        "heading" in style_name
        or style_name.startswith("заголовок")
        or para.text.isupper()
        and len(para.text.split()) <= 6
    )


def process_docx(file_content: bytes, title: str, author: str, group: int, subject: str) -> Dict[str, Any]:
    doc = Document(BytesIO(file_content))

    parts: List[Dict] = []
    current_heading = "основная часть"
    current_texts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if is_heading(para) or (len(text) < 80 and text.endswith((".", ":")) is False and text.isupper()):
            if current_texts:
                parts.append({
                    "type": detect_section_type(current_heading),
                    "chunks": split_into_chunks(" ".join(current_texts)),
                })
            current_heading = text
            current_texts = []
        else:
            current_texts.append(text)

    if current_texts:
        parts.append({
            "type": detect_section_type(current_heading),
            "chunks": split_into_chunks(" ".join(current_texts)),
        })

    if not parts:
        all_text = " ".join(p.text for p in doc.paragraphs if p.text.strip())
        if all_text:
            parts = [{"type": "основная часть", "chunks": split_into_chunks(all_text)}]

    all_text = " ".join(c["text"] for p in parts for c in p["chunks"])
    words_count = len(all_text.split())
    flesh_index = compute_flesh_index(all_text)
    keyword_density = compute_keyword_density(all_text)

    part_types = [p["type"] for p in parts]
    introduction = any("введен" in t or "вводн" in t for t in part_types)
    conclusion = any("заключен" in t or "вывод" in t for t in part_types)
    bibliography = any("литератур" in t or "библиограф" in t or "список" in t for t in part_types)

    return {
        "title": title,
        "author": author,
        "group": group,
        "subject": subject,
        "parts": parts,
        "words_count": words_count,
        "flesh_index": flesh_index,
        "keyword_density": keyword_density,
        "introduction": introduction,
        "conclusion": conclusion,
        "bibliography": bibliography,
    }
